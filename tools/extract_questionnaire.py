import json
import sys
from pathlib import Path

from docx import Document


def cell_text(cell):
    parts = []
    for paragraph in cell.paragraphs:
        text = " ".join(run.text for run in paragraph.runs).strip()
        if text:
            parts.append(text)
    return "\n".join(parts).strip()


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: extract_questionnaire.py input.docx output.json")

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    doc = Document(input_path)

    blocks = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(
                {
                    "kind": "paragraph",
                    "style": paragraph.style.name if paragraph.style else "",
                    "text": text,
                }
            )

    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell_text(cell) for cell in row.cells])
        blocks.append({"kind": "table", "table_index": table_index, "rows": rows})

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(blocks, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
